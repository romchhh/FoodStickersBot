import openpyxl
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def set_margins(doc, top=0.5, right=0, bottom=0, left=0):
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(top)
        section.right_margin = Cm(right)
        section.bottom_margin = Cm(bottom)
        section.left_margin = Cm(left)

def read_orders_from_excel(file_path):
    names = []
    orders = []
    try:
        workbook = openpyxl.load_workbook(file_path)
        sheet = workbook.active
        for row in sheet.iter_rows(min_row=2, max_col=1, values_only=True):
            data = row[0]

            if data is not None:
                cleaned_data = data.replace("<b>", "").replace("</b>", "")

                split_data = cleaned_data.split(':')
                if len(split_data) == 2:
                    name, order_info = split_data
                    order_split = order_info.split('(')
                    if order_split:
                        order = order_split[0].strip()
                        names.append(name.strip())
                        orders.append(order)

    except Exception as e:
        print(f"Error processing the file. Please try again.")
    return names, orders

def generate_word_from_excel_orders(input_file, output_file):
    try:
        document = Document()

        set_margins(document, 1.16, 0, 0, 0)
        names, orders = read_orders_from_excel(input_file)
        table = document.add_table(rows=1, cols=3)
        index_to_delete = 0
        row = table.rows[index_to_delete]
        row._element.getparent().remove(row._element)

        column_widths = [7.6, 7.2, 7.2]  
        for i, width in enumerate(column_widths):
            table.columns[i].width = Cm(width)

        for row in table.rows:
            for cell in row.cells:
                cell._element.clear_content()

        for i, (name, order) in enumerate(zip(names, orders)):
            if name:  
                if i % 3 == 0: 
                    cells = table.add_row().cells
                for row in table.rows:
                    row.height = Cm(3.51)

                cell = cells[i % 3]
                paragraph = cell.paragraphs[0]
                run_name = paragraph.add_run(name)
                run_name.bold = True
                run_name.font.size = Pt(11)
                run_name.font.name = 'Times New Roman'

                paragraph.add_run('\n') 

                order_lines = [line.strip() for line in order.split(',')]
                for line in order_lines:
                    run_order = paragraph.add_run(line)
                    run_order.font.size = Pt(10)
                    run_order.font.name = 'Times New Roman'
                    paragraph.add_run('\n')  

                cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        document.save(output_file)

    except Exception:
        raise Exception(f"Error generating stickers. Please try again.")
