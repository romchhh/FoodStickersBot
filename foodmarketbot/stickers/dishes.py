import openpyxl
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import re
from docx.shared import Pt, Cm

def set_margins(doc, top=0.5, right=0.5, bottom=0.5, left=0.5):
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(top)
        section.right_margin = Cm(right)
        section.bottom_margin = Cm(bottom)
        section.left_margin = Cm(left)


         
def process_dish_name(dish_name):
    if len(dish_name) > 21:
        words = dish_name.split()

        new_dish_name = ""
        current_line_length = 0

        for word in words:
            if current_line_length + len(word) <= 21:
                new_dish_name += word + " "
                current_line_length += len(word) + 1
            else:
                new_dish_name += "\n" + word + " "
                current_line_length = len(word) + 1

        return new_dish_name.strip()
    else:
        return dish_name + '\n'


def read_dishes_from_excel(file_path):
    dishes = []
    try:
        workbook = openpyxl.load_workbook(file_path)
        sheet_names = workbook.sheetnames

        first_sheet_name = sheet_names[0]
        sheet = workbook[first_sheet_name]

        max_row = sheet.max_row

        for row_number in range(2, max_row + 1):
            row = sheet[row_number]

            if len(row) >= 4:
                dish_info = row[1].value 
                dish_count = row[3].value  
                
                if dish_info and dish_count:
                    match = re.match(r'(.+?)(?:\((.*?)\))?$', dish_info)
                    if match:
                        dish_name, dish_description = match.groups()
                        dishes.append((dish_name.strip(), dish_description.strip() if dish_description else "", dish_count))
            else:
                print(f"Row {row_number} has fewer than 4 elements:", row)

    except FileNotFoundError:
        print(f"Error: File not found - {file_path}")
    except openpyxl.utils.exceptions.InvalidFileException:
        print(f"Error: Invalid Excel file - {file_path}")
    except Exception as e:
        print(f"Помилка при читанні страв з Excel, спробуйте будь ласка ще раз.")
    
    return dishes

def generate_word_from_excel_dishes(input_file, output_file, selected_date_dishes):
    try:
        document = Document()

        set_margins(document, top=0.5, right=0.5, bottom=1.5, left=0.85)

        dishes = read_dishes_from_excel(input_file)

        table = document.add_table(rows=1, cols=5)
        
        for col in table.columns:
            col.width = Cm(4)

        cell_counter = 0

        for dish_name, dish_description, dish_count in dishes:
            for _ in range(int(dish_count)):
                row_idx = cell_counter // 5
                col_idx = cell_counter % 5

                    
                for row in table.rows:
                    row.height = Cm(2.1)
                    
                while len(table.rows) <= row_idx:
                    table.add_row()

                cell = table.cell(row_idx, col_idx)
                processed_dish_name = process_dish_name(dish_name)
                cell.text = f"    \n\n{processed_dish_name}    \n Виготовлено: {selected_date_dishes.strftime('%d.%m')}"
                    
                run = cell.paragraphs[0].runs[0]
                run.font.size = Pt(7.5)

                cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                cell_counter += 1

        document.save(output_file)

    except Exception as e:
        print(f"Помилка генерації наліпок, спробуйте будь ласка ще раз.")


         
def process_dish_name(dish_name):
    if len(dish_name) > 21:
        words = dish_name.split()

        new_dish_name = ""
        current_line_length = 0

        for word in words:
            if current_line_length + len(word) <= 21:
                new_dish_name += word + " "
                current_line_length += len(word) + 1
            else:
                new_dish_name += "\n" + word + " "
                current_line_length = len(word) + 1

        return new_dish_name.strip()
    else:
        return dish_name + '\n'


