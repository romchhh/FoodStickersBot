import openpyxl
import re
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def set_margins(doc, top=0.5, right=0.5, bottom=0.5, left=0.5):
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(top)
        section.right_margin = Cm(right)
        section.bottom_margin = Cm(bottom)
        section.left_margin = Cm(left)

def get_complex_complex(file_path):
    pattern = r"^(?=.*[a-zA-Z])\w*$"
    complex_dishes = {}

    try:
        workbook = openpyxl.load_workbook(file_path)
        sheet = workbook[workbook.sheetnames[0]]
        max_row = sheet.max_row

        current_complex = None  
        stop_adding = False  

        for index in range(2, max_row + 1):
            row = sheet[index]

            if index < 5:
                continue

            if (
                (re.match(pattern, str(row[0].value)) and sheet[index+1][4].value is not None)
                or
                (isinstance(row[0].value, str) and (
                    "complex" in row[0].value.lower() or
                    "mini" in row[0].value.lower()
                ))
            ):
                current_complex = row[0].value 
                stop_adding = False  
                if current_complex not in complex_dishes:
                    complex_dishes[current_complex] = []

            if current_complex is not None and row[1].value is not None and not stop_adding:
                if "столові прибори" not in str(row[1].value).lower():
                    quantity = int(row[4].value) if row[4].value is not None else 1
                    complex_dishes[current_complex].extend([[row[1].value, row[3].value, row[4].value]] * quantity)
                    
            elif isinstance(row[0].value, int) or (isinstance(row[0].value, str) and not any(keyword in row[0].value.lower() for keyword in ["complex", "mini"])):
                stop_adding = True

        return complex_dishes

    except Exception as e:
        print(f"Помилка генерації наліпок, спробуйте будь ласка ще раз.")

def update_excel_values(file_path, complex1_value, complex2_value, mini1_value, mini2_value):
    try:
        workbook = openpyxl.load_workbook(file_path)
        sheet = workbook[workbook.sheetnames[0]]
        max_row = sheet.max_row

        values_to_replace = [complex1_value, complex2_value, mini1_value, mini2_value]
        replace_index = 0

        for index in range(2, max_row + 1):
            row = sheet[index]

            if isinstance(row[4].value, (int, float)) and row[4].value != 0 and replace_index < len(values_to_replace):
                row[4].value = values_to_replace[replace_index]
                replace_index += 1

        workbook.save(file_path)

    except Exception as e:
        print(f"Помилка генерації наліпок, спробуйте будь ласка ще раз.")
        
        
def generate_word_from_excel_complex(input_file, output_file, complex1_value, complex2_value, mini1_value, mini2_value):
    try:
        document = Document()
        set_margins(document, 0.33, 0.3, 0, 0.48)

        update_excel_values(input_file, complex1_value, complex2_value, mini1_value, mini2_value)

        complex_dishes = get_complex_complex(input_file)
        table = document.add_table(rows=1, cols=3)
        column_widths = [6.69, 8.3, 6]  
        for i, width in enumerate(column_widths):
            table.columns[i].width = Cm(width)

        cell_counter = 0

        for complex_name, dish_list in complex_dishes.items():
            dish_names = dict.fromkeys(dish[0].split('(')[0] if dish[0] else None for dish in dish_list)
            dish_names = list(dish_names.keys())
            dish_quantity = [dish[2] if dish[2] else None for dish in dish_list]

            if not any(dish_quantity):
                continue

            for quantity in dish_quantity:
                if quantity is None:
                    continue

                row_idx = cell_counter // 3
                col_idx = cell_counter % 3
                
                for row in table.rows:
                    row.height = Cm(3.59)

                while len(table.rows) <= row_idx:
                    table.add_row()

                cell = table.cell(row_idx, col_idx)
                p = cell.add_paragraph()
                run = p.add_run(f"{complex_name}\n")
                run.font.bold = True
                run = p.add_run("\n".join(dish_name for dish_name in dish_names))
                run.font.bold = False

                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.space_before = Pt(0)

                for run in p.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(8)

                cell_counter += 1

        document.save(output_file)

    except Exception as e:
        print(f"Помилка генерації наліпок, спробуйте будь ласка ще раз.")

