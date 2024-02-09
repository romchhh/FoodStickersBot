import openpyxl
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from googletrans import Translator
import re

def translate_to_english(text):
    translator = Translator()
    translation = translator.translate(text, dest='en', src='uk')
    translation.text = translation.text.replace("Kiev-style cutlet", "Chiken Kyiv")
    return translation.text

def set_margins(doc, top=0.5, right=0.5, bottom=0.5, left=0.5):
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(top)
        section.right_margin = Cm(right)
        section.bottom_margin = Cm(bottom)
        section.left_margin = Cm(left)
        
def get_complex_dishes(file_path):
    pattern = r"^(?=.*[a-zA-Z])\w*$"
    complex_dishes = {}
    another_dishes = []

    try:
        workbook = openpyxl.load_workbook(file_path)
        sheet = workbook[workbook.sheetnames[0]]
        max_row = sheet.max_row

        current_complex = None  
        stop_adding = False  

        for index in range(2, max_row + 1):
            row = sheet[index]

            if index < 5:
                continue

            if (
                (re.match(pattern, str(row[0].value)) and sheet[index+1][4].value is not None)
                or
                (isinstance(row[0].value, str) and (
                    "complex" in row[0].value.lower() or
                    "mini" in row[0].value.lower()
                ))
            ):
                current_complex = row[0].value 
                stop_adding = False 
                if current_complex not in complex_dishes:
                    complex_dishes[current_complex] = []

            if current_complex is not None and row[1].value is not None and not stop_adding:
                if "столові прибори" not in str(row[1].value).lower():
                    quantity = int(row[4].value) if row[4].value is not None else 1
                    complex_dishes[current_complex].extend([[row[1].value, row[3].value, row[4].value]] * quantity)
                    print(quantity)
            elif isinstance(row[0].value, int) or (isinstance(row[0].value, str) and not any(keyword in row[0].value.lower() for keyword in ["complex", "mini"])):
                # If a new section starts, stop adding entries to the current complex
                stop_adding = True
                another_dishes.append([row[1].value, row[3].value, row[4].value])

        return complex_dishes, another_dishes

    except Exception as e:
        print(f"Помилка генерації наліпок, спробуйте будь ласка ще раз.")

        
def read_fridge_from_excel(file_path):
    dishes = []
    try:
        workbook = openpyxl.load_workbook(file_path)
        sheet_names = workbook.sheetnames

        first_sheet_name = sheet_names[0]
        sheet = workbook[first_sheet_name]

        max_row = sheet.max_row

        for row_number in range(2, max_row + 1):
            row = sheet[row_number]

            if len(row) >= 4:
                dish_type = str(row[0].value)
                dish_info = str(row[1].value)
                dish_count = str(row[4].value)  
                dish_value = str(row[3].value)

                if dish_info and dish_count:
                    match = re.match(r'([^(]+)\(([^)]+)\)', dish_info)
                    if match:
                        dish_name, dish_description = match.groups()
                        dishes.append((dish_name.strip(), dish_description.strip(), dish_type.strip(), dish_count, dish_value))
            else:
                print(f"Row {row_number} has fewer than 4 elements:", row)

    except Exception as e:
        print(f"Помилка при читанні страв з Excel, спробуйте будь ласка ще раз.")
        
    return dishes

def generate_word_from_excel_fridge(input_file, output_file, selected_date_fridge):
    try:
        document = Document()
        set_margins(document, 1.7, 0.5, 0.0, 1.2)

        complex_dishes, another_dishes = get_complex_dishes(input_file)

        table = document.add_table(rows=1, cols=3)
        index_to_delete = 0
        row = table.rows[index_to_delete]
        row._element.getparent().remove(row._element)

        column_widths = [8.3, 3.5, 8.2]
        for i, width in enumerate(column_widths):
            table.columns[i].width = Cm(width)

        current_column = 0
        i = 0
        

        for complex_name, dish_list in complex_dishes.items():
            dish_names = dict.fromkeys(dish[0].split('(')[0] if dish[0] else None for dish in dish_list)
            dish_names = list(dish_names.keys())
            
            dish_price = [dish[1] if dish[1] else None for dish in dish_list]
            dish_quantity = [dish[2] if dish[2] else None for dish in dish_list]

            if not any(dish_quantity):
                continue

            transled_dish_names = list()

            for index in range(0, len(dish_names)):
                
                translated_dish = translate_to_english(dish_names[index])
                transled_dish_names.append(translated_dish)
                if  index == dish_names:
                    transled_dish_names = transled_dish_names[::-1]
        
            for quantity in dish_quantity:
                if quantity is None:
                    continue

                if current_column == 0:
                    sixth_row_height = [4.23, 4.87, 4.97, 4.74, 4.23, 2.74]
                    new_row = table.add_row()
                    new_row.height = Cm(sixth_row_height[i])
                    i += 1
                    if i == 6:
                        i = 0

                cell = new_row.cells[current_column]

                if not any(cell.text.strip()):
                    name_paragraph = cell.paragraphs[0]
                    description_paragraph = cell.add_paragraph()
                    bold_text_paragraph = cell.add_paragraph()
                else:
                    name_paragraph = cell.add_paragraph()
                    description_paragraph = cell.add_paragraph()
                    bold_text_paragraph = cell.add_paragraph()

                name_run = name_paragraph.add_run(f"{complex_name}" if complex_name else "")
                name_run.bold = True

                name_paragraph.paragraph_format.space_before = Pt(0)

                description_text = "\n".join(
                    f"{name}\n{translation}" if name and translation else ""
                    for name, translation in zip(dish_names, transled_dish_names)
                )
                description_run = description_paragraph.add_run(description_text)

                bold_text_run = bold_text_paragraph.add_run(
                    f"Вжити до {selected_date_fridge.strftime('%d.%m')}                              "
                    f"Ціна {dish_price[0]} грн" if dish_price[0] else ""
                )
                bold_text_run.bold = True

                for p in [name_paragraph, description_paragraph, bold_text_paragraph]:
                    for run in p.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(8)

                for p in [name_paragraph, description_paragraph, bold_text_paragraph]:
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    p.paragraph_format.space_after = Pt(0)
                    p.paragraph_format.space_before = Pt(0)

                current_column = 2 if current_column == 0 else 0

        for dish_with, dish_value, dish_count in another_dishes:
            dish_name = dish_with.split('(')[0] if dish_with else None
            dish_parts = dish_with.split('(') if dish_with else []
            translated_dish_name = translate_to_english(dish_name) if dish_name else None

            if translated_dish_name is not None:
                translated_description = translate_to_english(dish_parts[1]) if len(dish_parts) >= 2 and dish_parts[1] else None

                for _ in range(int(dish_count)) if dish_count is not None else range(0):
                    if current_column == 0:
                        sixth_row_height = [4.23, 4.87, 4.97, 4.74, 4.23, 2.74]
                        new_row = table.add_row()
                        new_row.height = Cm(sixth_row_height[i])
                        i += 1
                        if i == 6:
                            i = 0

                    cell = new_row.cells[current_column]

                    if not any(cell.text.strip()):
                        name_paragraph = cell.paragraphs[0]
                        description_paragraph = cell.add_paragraph()
                        bold_text_paragraph = cell.add_paragraph()
                    else:
                        name_paragraph = cell.add_paragraph()
                        description_paragraph = cell.add_paragraph()
                        bold_text_paragraph = cell.add_paragraph()

                    name_run = name_paragraph.add_run(f"{dish_name}\\{translated_dish_name}".replace('(', '').replace(')', '') if dish_name else "")
                    name_run.bold = True

                    name_paragraph.paragraph_format.space_before = Pt(0)

                    if translated_description is not None:
                        description_run = description_paragraph.add_run(f"{dish_parts[1]}".replace('(', '').replace(')', '') + "\n" + f"{translated_description}".replace('(', '').replace(')', ''))
                    else:
                        description_run = description_paragraph.add_run(f"{dish_parts[1]}\n".replace('(', '').replace(')', '') if len(dish_parts) >= 2 else "")

                    bold_text_run = bold_text_paragraph.add_run(
                        f"Вжити до {selected_date_fridge.strftime('%d.%m')}                              "
                        f"Ціна {dish_value} грн" if dish_value else ""
                    )
                    bold_text_run.bold = True

                    for p in [name_paragraph, description_paragraph, bold_text_paragraph]:
                        for run in p.runs:
                            run.font.name = 'Times New Roman'
                            run.font.size = Pt(8)

                    for p in [name_paragraph, description_paragraph, bold_text_paragraph]:
                        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        p.paragraph_format.space_after = Pt(0)
                        p.paragraph_format.space_before = Pt(0)

                    current_column = 2 if current_column == 0 else 0

        document.save(output_file)

    except Exception as e:
        print(f"Помилка генерації наліпок, спробуйте будь ласка ще раз.")