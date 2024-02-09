import openpyxl
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def set_margins(doc, top=0.5, right=0.5, bottom=0.5, left=0.5):
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(top)
        section.right_margin = Cm(right)
        section.bottom_margin = Cm(bottom)
        section.left_margin = Cm(left)

def read_names_from_excel(file_path):
    names = []
    try:
        workbook = openpyxl.load_workbook(file_path)
        sheet = workbook.active
        for row in sheet.iter_rows(min_row=2, max_col=2, values_only=True):
            name = row[1]  
            if name: 
                names.append(name)
    except Exception as e:
        print(f"Помилка при читанні імен з Excel, спробуйте будь ласка ще раз.")

    return names


def generate_word_from_excel_names(input_file, output_file):
    try:
        document = Document()

        set_margins(document, 1.3, 0, 0.5, 0)

        names = read_names_from_excel(input_file)

        table = document.add_table(rows=1, cols=3)
        index_to_delete = 0
        row = table.rows[index_to_delete]
        row._element.getparent().remove(row._element)
        
        column_widths = [6.66, 8.5, 6]  
        for i, width in enumerate(column_widths):
            table.columns[i].width = Cm(width)
    
        for row in table.rows:
            for cell in row.cells:
                cell._element.clear_content()

        for i, name in enumerate(names):
            if name:  
                if i % 3 == 0: 
                    cells = table.add_row().cells
                for row in table.rows:
                    row.height = Cm(3.7)

                cell = cells[i % 3]
                paragraph = cell.paragraphs[0]
                run = paragraph.add_run(name + '\n' + "гарного дня)")
                run.bold = True
                run.font.size = Pt(11)
                run.font.name = 'Times New Roman'

                cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        document.save(output_file)

    except Exception as e:
        raise Exception(f"Помилка генерації наліпок, спробуйте будь ласка ще раз.")  
    